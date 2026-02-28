"""Service for extracting text from document files (PDF, DOCX, XLSX, CSV)."""

import asyncio
import csv
import io
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

MAX_VAULT_CHARS = 50_000
MAX_SHEETS = 10


class DocumentExtractor:
    """Extract text content from document files."""

    async def extract(self, file_bytes: bytes, file_name: str) -> str:
        ext = Path(file_name).suffix.lower()

        extractors = {
            ".pdf": self._extract_pdf,
            ".docx": self._extract_docx,
            ".xlsx": self._extract_xlsx,
            ".csv": self._extract_csv,
        }

        extractor = extractors.get(ext)
        if not extractor:
            raise ValueError(f"Unsupported format: {ext}")

        text = await asyncio.to_thread(extractor, file_bytes)
        return text

    def _extract_pdf(self, file_bytes: bytes) -> str:
        import fitz

        pages = []
        try:
            doc = fitz.open(stream=file_bytes, filetype="pdf")
        except Exception as e:
            if "password" in str(e).lower() or "encrypted" in str(e).lower():
                raise ValueError("Документ защищён паролем") from e
            raise

        try:
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text().strip()
                if text:
                    pages.append(text)
        finally:
            doc.close()

        if not pages:
            raise ValueError(
                "PDF содержит только изображения, текст не найден. "
                "Попробуй сначала распознать текст через OCR."
            )

        return "\n\n---\n\n".join(pages)

    def _extract_docx(self, file_bytes: bytes) -> str:
        from docx import Document

        doc = Document(io.BytesIO(file_bytes))
        parts: list[str] = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)

        for table in doc.tables:
            table_lines = self._table_to_markdown(
                [[cell.text.strip() for cell in row.cells] for row in table.rows]
            )
            if table_lines:
                parts.append(table_lines)

        return "\n\n".join(parts)

    def _extract_xlsx(self, file_bytes: bytes) -> str:
        from openpyxl import load_workbook

        wb = load_workbook(
            io.BytesIO(file_bytes), read_only=True, data_only=True
        )

        parts: list[str] = []
        sheet_names = wb.sheetnames[:MAX_SHEETS]

        for sheet_name in sheet_names:
            ws = wb[sheet_name]
            rows: list[list[str]] = []

            for row in ws.iter_rows(values_only=True):
                str_row = [str(cell) if cell is not None else "" for cell in row]
                if any(c for c in str_row):
                    rows.append(str_row)

            if rows:
                header = f"### {sheet_name}" if len(sheet_names) > 1 else ""
                table_md = self._table_to_markdown(rows)
                if header:
                    parts.append(f"{header}\n\n{table_md}")
                else:
                    parts.append(table_md)

        wb.close()

        if len(wb.sheetnames) > MAX_SHEETS:
            parts.append(
                f"\n*[Показаны первые {MAX_SHEETS} листов из {len(wb.sheetnames)}]*"
            )

        return "\n\n".join(parts)

    def _extract_csv(self, file_bytes: bytes) -> str:
        text = None
        for encoding in ("utf-8", "cp1251", "latin-1"):
            try:
                text = file_bytes.decode(encoding)
                break
            except (UnicodeDecodeError, ValueError):
                continue

        if text is None:
            raise ValueError("Не удалось определить кодировку CSV файла")

        reader = csv.reader(io.StringIO(text))
        rows: list[list[str]] = []
        for row in reader:
            if any(cell.strip() for cell in row):
                rows.append([cell.strip() for cell in row])

        if not rows:
            raise ValueError("CSV файл пуст")

        return self._table_to_markdown(rows)

    @staticmethod
    def _table_to_markdown(rows: list[list[str]]) -> str:
        if not rows:
            return ""

        max_cols = max(len(row) for row in rows)
        normalized = [row + [""] * (max_cols - len(row)) for row in rows]

        header = normalized[0]
        lines = ["| " + " | ".join(header) + " |"]
        lines.append("| " + " | ".join(["---"] * max_cols) + " |")

        for row in normalized[1:]:
            lines.append("| " + " | ".join(row) + " |")

        return "\n".join(lines)

    @staticmethod
    def truncate_for_vault(text: str, attachment_path: str) -> str:
        if len(text) <= MAX_VAULT_CHARS:
            return text

        truncated = text[:MAX_VAULT_CHARS]
        last_para = truncated.rfind("\n\n")
        if last_para > MAX_VAULT_CHARS * 0.8:
            truncated = truncated[:last_para]
        else:
            last_sentence = truncated.rfind(". ")
            if last_sentence > MAX_VAULT_CHARS * 0.8:
                truncated = truncated[: last_sentence + 1]

        truncated += (
            f"\n\n---\n*[Текст обрезан ({len(text)} символов). "
            f"Полный документ: {attachment_path}]*"
        )
        return truncated
